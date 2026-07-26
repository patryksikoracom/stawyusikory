from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
FULL_PATH = OUTPUT_DIR / "SOP_OCHRONA_MALOLETNICH_STAWY_U_SIKORY_V0.9_PROJEKT.docx"
CHILD_PATH = OUTPUT_DIR / "SOP_OCHRONA_MALOLETNICH_DLA_DZIECI_V0.9_PROJEKT.docx"

INK = "19352D"
GREEN = "315D4C"
LIGHT_GREEN = "E5EFE8"
LAKE = "DCEBEC"
MUTED = "65736C"
GOLD = "8A651B"
LIGHT_GOLD = "F7EDCF"
RED = "8F402C"
LIGHT_RED = "FCEAE4"
GRAY = "F2F4F3"
WHITE = "FFFFFF"
BLACK = "111111"
TABLE_WIDTH = 9360


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float = 11,
    color: str = BLACK,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120):
    assert sum(widths) == TABLE_WIDTH
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_repeat_table_header(row):
    mark_header_row(row)


def set_table_borders(table, color="D7DEDA", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_paragraph_border(paragraph, color=GREEN, size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_paragraph_box(paragraph, *, color: str, fill: str, size: str = "8"):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), color)
        borders.append(border)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Strona ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([color, underline])
    run.append(run_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_numbering_definition(document: Document, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def configure_document(document: Document, *, child: bool = False):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial" if child else "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), normal.font.name)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), normal.font.name)
    normal.font.size = Pt(13.5 if child else 11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15 if child else 1.25

    heading_tokens = {
        "Heading 1": (21 if child else 16, GREEN, 16 if child else 18, 8 if child else 10),
        "Heading 2": (16 if child else 13, GREEN, 12 if child else 14, 6 if child else 7),
        "Heading 3": (14.5 if child else 12, INK, 9 if child else 10, 4 if child else 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = normal.font.name
        style._element.rPr.rFonts.set(qn("w:ascii"), normal.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), normal.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(
        "STAWY U SIKORY  |  OCHRONA MAŁOLETNICH"
        if not child
        else "STAWY U SIKORY  |  TWOJE BEZPIECZEŃSTWO"
    )
    set_run_font(run, name=normal.font.name, size=8.5, color=MUTED, bold=True)
    set_paragraph_border(header, color="CDD8D3", size="4")
    add_page_field(section.footer.paragraphs[0])


def add_title_page(document: Document, *, child: bool = False):
    for _ in range(4 if child else 6):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("WERSJA DLA DZIECI I MŁODZIEŻY" if child else "STANDARD OPERACYJNY • WERSJA 0.9")
    set_run_font(run, name="Arial" if child else "Calibri", size=11, color=GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(16)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(
        "Tu masz prawo czuć się bezpiecznie"
        if child
        else "Standardy ochrony małoletnich"
    )
    set_run_font(
        run,
        name="Arial" if child else "Calibri",
        size=28 if child else 30,
        color=INK,
        bold=True,
    )
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run(
        "Krótka informacja dla osób poniżej 18 lat"
        if child
        else "Stawy u Sikory • obiekt noclegowy • Nowa Wieś 95, 66-350 Bledzew"
    )
    set_run_font(
        run,
        name="Arial" if child else "Calibri",
        size=15,
        color=GREEN,
        bold=not child,
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_box(p, color=RED, fill=LIGHT_RED, size="10")
    run = p.add_run(
        "PROJEKT DO ZATWIERDZENIA — NIE OBOWIĄZUJE"
        if not child
        else "JEŚLI GROZI CI NIEBEZPIECZEŃSTWO, DZWOŃ 112"
    )
    set_run_font(run, name="Arial" if child else "Calibri", size=13 if child else 12, color=RED, bold=True)
    document.add_paragraph()
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Opracowano: 26 lipca 2026 r. • przegląd źródeł: stan na 26 lipca 2026 r."
        if not child
        else "Możesz też poprosić o pomoc pracownika albo zadzwonić pod 116 111."
    )
    set_run_font(run, name="Arial" if child else "Calibri", size=11 if child else 10, color=MUTED)
    document.add_page_break()


def add_callout(
    document: Document,
    title: str,
    body: str,
    *,
    tone: str = "info",
    child: bool = False,
):
    color, fill = {
        "info": (GREEN, LIGHT_GREEN),
        "warn": (GOLD, LIGHT_GOLD),
        "danger": (RED, LIGHT_RED),
    }[tone]
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    set_paragraph_box(p, color=color, fill=fill, size="8")
    run = p.add_run(title)
    set_run_font(
        run,
        name="Arial" if child else "Calibri",
        size=13 if child else 11,
        color=color,
        bold=True,
    )
    run.add_break()
    run = p.add_run(body)
    set_run_font(
        run,
        name="Arial" if child else "Calibri",
        size=12.5 if child else 10.5,
        color=BLACK,
    )


def add_bullets(document: Document, items: Iterable[str], bullet_num_id: int, *, child=False):
    for item in items:
        p = document.add_paragraph()
        apply_num(p, bullet_num_id)
        run = p.add_run(item)
        set_run_font(run, name="Arial" if child else "Calibri", size=13.5 if child else 11)


def add_steps(document: Document, items: Iterable[str], number_num_id: int, *, child=False):
    local_num_id = add_numbering_definition(document, bullet=False)
    for item in items:
        p = document.add_paragraph()
        apply_num(p, local_num_id)
        run = p.add_run(item)
        set_run_font(run, name="Arial" if child else "Calibri", size=13.5 if child else 11)


def add_key_value_table(document: Document, rows: list[tuple[str, str]]):
    table = document.add_table(rows=1, cols=2)
    for idx, text in enumerate(("Pole", "Treść")):
        set_cell_shading(table.rows[0].cells[idx], GREEN)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=9.5, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], GRAY)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=10, color=INK, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=10.5, color=BLACK)
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)
    document.add_paragraph()


def add_roles_table(document: Document):
    table = document.add_table(rows=1, cols=3)
    headers = ("Rola", "Odpowiedzialność", "Nie może")
    for idx, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], GREEN)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=10, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    rows = [
        (
            "Właściciel podmiotu",
            "zatwierdza standard, powołuje koordynatora i zastępcę, zatwierdza retencję, zasoby i przeglądy",
            "uznać projektu 0.9 za obowiązujący bez uzupełnienia decyzji i przeszkolenia osób",
        ),
        (
            "Koordynator",
            "przyjmuje eskalacje, decyduje o zawiadomieniach, prowadzi rejestr, zabezpiecza dowody i organizuje wsparcie",
            "prowadzić prywatnego śledztwa, konfrontować dziecka z osobą podejrzewaną lub obiecywać tajemnicy",
        ),
        (
            "Osoba wydająca klucz/kod",
            "wykonuje sześć kroków procedury, nie wydaje dostępu przy niewyjaśnionej sytuacji, eskaluje bez zwłoki",
            "zapisywać danych dziecka w Stawy OS ani kopiować dokumentów",
        ),
        (
            "Każdy członek personelu",
            "reaguje na sygnały, zapewnia bezpieczeństwo i przekazuje fakty koordynatorowi albo służbom",
            "ignorować zgłoszenia, prowadzić przesłuchania lub rozpowszechniać informacje",
        ),
    ]
    for role, responsibility, forbidden in rows:
        cells = table.add_row().cells
        values = (role, responsibility, forbidden)
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=9.5, color=BLACK, bold=idx == 0)
    set_table_geometry(table, [1800, 4320, 3240])
    set_table_borders(table)
    document.add_paragraph()


def add_contact_table(document: Document):
    table = document.add_table(rows=1, cols=3)
    for idx, text in enumerate(("Kiedy", "Kontakt", "Dane")):
        set_cell_shading(table.rows[0].cells[idx], GREEN)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=10, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("natychmiastowe zagrożenie życia/zdrowia", "Numer alarmowy", "112"),
        ("pilne zgłoszenie Policji", "Dyżurny KPP Międzyrzecz", "47 79 212 11 / 47 79 212 12"),
        ("zawiadomienie sądu opiekuńczego", "SR Międzyrzecz, III Wydział Rodzinny i Nieletnich", "95 742 40 03 • rodzinny@miedzyrzecz.sr.gov.pl"),
        ("pomoc społeczna / sytuacja rodzinna", "GOPS Bledzew", "95 743 66 23 • gops@bledzew.pl"),
        ("wsparcie dziecka 24/7", "Telefon zaufania", "116 111"),
        ("wsparcie dziecka 24/7", "Rzecznik Praw Dziecka", "800 12 12 12 • czat.brpd.gov.pl"),
        ("kontakt obiektu — noclegi", "Marcin", "501 510 005 • marcin@stawyusikory.pl"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=9.5, color=BLACK, bold=idx == 0)
    set_table_geometry(table, [2740, 3380, 3240])
    set_table_borders(table)
    document.add_paragraph()


def add_sources(document: Document):
    document.add_heading("Załącznik G. Podstawa i źródła", level=1)
    p = document.add_paragraph(
        "Źródła sprawdzono 26 lipca 2026 r. Dokument nie zastępuje indywidualnej porady prawnej ani oceny skutków dla ochrony danych."
    )
    p.paragraph_format.space_after = Pt(10)
    sources = [
        (
            "Ustawa — tekst jednolity, Dz.U. 2026 poz. 110, szczególnie art. 22c ust. 3–7",
            "https://eli.gov.pl/api/acts/DU/2026/110/text/U/D20260110Lj.pdf",
        ),
        (
            "Kodeks postępowania cywilnego — tekst jednolity, Dz.U. 2026 poz. 468, art. 572",
            "https://eli.gov.pl/api/acts/DU/2026/468/text.pdf",
        ),
        (
            "Ministerstwo Sprawiedliwości — wytyczne do standardów ochrony dzieci",
            "https://www.gov.pl/web/sprawiedliwosc/standardy-ochrony-maloletnich---wytyczne.",
        ),
        (
            "UODO — jak stosować ustawę z poszanowaniem minimalizacji i bezpieczeństwa danych",
            "https://uodo.gov.pl/pl/138/3278",
        ),
        (
            "KSSiP — publiczny przykład standardów dla obiektów świadczących usługi hotelarskie",
            "https://www.kssip.gov.pl/sites/default/files/zalacznik_nr_1_standardy_ochrony_maloletnich_w_obiektach_kssip.pdf",
        ),
        (
            "Fundacja Dajemy Dzieciom Siłę — materiały dla branży hotelarskiej",
            "https://standardy.fdds.pl/blog/15/webinar-wprowadzanie-standardow-ochrony-dzieci-w-obiektach-swiadczacych-uslugi-hotelarskie",
        ),
        (
            "KPP Międzyrzecz — aktualne dane kontaktowe",
            "https://miedzyrzecz.policja.gov.pl/go4/kontakt/dane-kontaktowe/47399,Dane-kontaktowe.html",
        ),
        (
            "Sąd Rejonowy w Międzyrzeczu — III Wydział Rodzinny i Nieletnich",
            "https://miedzyrzecz.sr.gov.pl/iii-wydzial-rodzinny-i-nieletnich,m,mg,2,219",
        ),
        (
            "Rzecznik Praw Dziecka — Dziecięcy Telefon Zaufania",
            "https://brpd.gov.pl/dzieciecy-telefon-zaufania-rzecznika-praw-dziecka/",
        ),
        (
            "Stawy u Sikory — publiczne dane kontaktowe obiektu",
            "https://stawyusikory.pl/kontakt-i-jak-dojechac/",
        ),
    ]
    for idx, (label, url) in enumerate(sources, start=1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(5)
        set_run_font(p.add_run(f"{idx}. {label}: "), size=9.5, color=BLACK)
        add_hyperlink(p, "otwórz źródło", url)


def build_full_document():
    doc = Document()
    configure_document(doc)
    bullet_id = add_numbering_definition(doc, bullet=True)
    number_id = add_numbering_definition(doc, bullet=False)
    add_title_page(doc)

    doc.add_heading("Karta dokumentu", level=1)
    add_key_value_table(
        doc,
        [
            ("Status", "Projekt 0.9 do zatwierdzenia — nie jest jeszcze podstawą wydania kluczy ani aktywacji w Stawy OS."),
            ("Podmiot i obiekt", "Podmiot prowadzący Stawy u Sikory, Nowa Wieś 95, 66-350 Bledzew. Formalną nazwę podmiotu należy potwierdzić w zarządzeniu zatwierdzającym."),
            ("Zakres", "Usługi noclegowe i turystyczne Stawy u Sikory oraz osoby wykonujące związane z nimi czynności."),
            ("Koordynator", "Do formalnego powołania funkcję organizacyjną pełni właściciel. Przed wejściem w życie trzeba imiennie powołać koordynatora i zastępcę."),
            ("Wersja / data", "0.9 / 26.07.2026"),
            ("Wejście w życie", "Wyłącznie data wskazana w podpisanym zarządzeniu po szkoleniu personelu, publikacji obu wersji i uruchomieniu bezpiecznego rejestru."),
            ("Przegląd", "Pierwszy przegląd po 6 miesiącach pilotażu; następnie co najmniej raz na 2 lata oraz po każdym poważnym zdarzeniu lub zmianie prawa."),
        ],
    )
    add_callout(
        doc,
        "CZTERY DECYZJE PRZED ZATWIERDZENIEM",
        "Właściciel musi: (1) potwierdzić formalną nazwę podmiotu, (2) powołać koordynatora i zastępcę, "
        "(3) zatwierdzić miejsce bezpiecznego rejestru oraz listę osób z dostępem, (4) zatwierdzić okresy retencji po konsultacji prawnej/RODO.",
        tone="warn",
    )

    doc.add_heading("Instrukcja alarmowa — pierwsze 10 minut", level=1)
    add_steps(
        doc,
        [
            "Jeżeli istnieje bezpośrednie zagrożenie życia lub zdrowia, najpierw dzwoń pod 112. Podaj lokalizację: Stawy u Sikory, Nowa Wieś 95, 66-350 Bledzew.",
            "Zapewnij dziecku bezpieczne, widoczne miejsce i spokojną obecność dorosłego pracownika. Nie zostawiaj go sam na sam z osobą, której dotyczy podejrzenie.",
            "Nie używaj siły i nie próbuj samodzielnie zatrzymywać kogokolwiek. Stosuj polecenia dyspozytora 112 lub Policji.",
            "Nie przesłuchuj dziecka. Wysłuchaj, nie oceniaj, nie obiecuj tajemnicy; zapisz później jego spontaniczne słowa możliwie dokładnie.",
            "Powiadom koordynatora. Jeżeli kontakt z nim opóźniałby ochronę dziecka, działaj bezpośrednio i poinformuj go po wezwaniu służb.",
            "Zabezpiecz dostępne dowody bez ich przeglądania, kopiowania lub rozpowszechniania ponad konieczny zakres; zachowaj oryginały i ciągłość dostępu.",
            "Utwórz kartę interwencji w bezpiecznym rejestrze. W Stawy OS zaznacz tylko „Wymaga reakcji” i później wpisz numer sprawy — bez danych dziecka i opisu zdarzenia.",
        ],
        number_id,
    )
    add_callout(
        doc,
        "Najważniejsza zasada",
        "Bezpieczeństwo dziecka ma pierwszeństwo przed komfortem obsługi, wydaniem klucza, płatnością i reputacją obiektu.",
        tone="danger",
    )
    doc.add_page_break()

    doc.add_heading("1. Cel, zakres i zasady nadrzędne", level=1)
    p = doc.add_paragraph(
        "Celem Standardów jest zapobieganie krzywdzeniu osób poniżej 18 lat, szybkie rozpoznawanie sygnałów zagrożenia "
        "oraz jednolite działanie personelu Stawy u Sikory. Standardy dotyczą pobytu, rezerwacji, wydania klucza lub kodu, "
        "kontaktu z personelem, usług dodatkowych i zdarzeń na terenie obiektu."
    )
    add_bullets(
        doc,
        [
            "Dobro i bezpieczeństwo dziecka są pierwszym kryterium decyzji.",
            "Każde dziecko jest traktowane z szacunkiem, bez dyskryminacji i z uwzględnieniem wieku, niepełnosprawności, neuroróżnorodności, języka i szczególnych potrzeb.",
            "Procedury stosuje się proporcjonalnie, spokojnie i bez automatycznego oskarżania gościa.",
            "Brak dokumentu dziecka sam w sobie nie dowodzi zagrożenia; wymaga spokojnego wyjaśnienia relacji innymi adekwatnymi metodami.",
            "Personel nie prowadzi śledztwa. Jego zadaniem jest bezpieczeństwo, zapis faktów i przekazanie sprawy właściwym służbom.",
            "Dane osobowe ogranicza się do minimum koniecznego do ochrony dziecka i wykonania obowiązków prawnych.",
        ],
        bullet_id,
    )

    doc.add_heading("2. Definicje robocze", level=1)
    add_key_value_table(
        doc,
        [
            ("Dziecko / małoletni", "Osoba, która nie ukończyła 18 lat."),
            ("Opiekun", "Rodzic, opiekun prawny lub inna osoba, której powierzono pieczę nad dzieckiem."),
            ("Krzywdzenie", "Przemoc fizyczna, psychiczna lub seksualna, zaniedbanie, wykorzystanie, handel ludźmi albo inne działanie lub zaniechanie zagrażające dobru dziecka."),
            ("Uzasadnione przypuszczenie", "Zestaw obserwacji, informacji lub niespójności, który racjonalnie wymaga ochrony i wyjaśnienia; nie oznacza przesądzenia winy."),
            ("Interwencja", "Działania służące natychmiastowemu bezpieczeństwu, eskalacji i zgłoszeniu właściwym instytucjom."),
            ("Stawy OS", "System operacyjny zapisujący wyłącznie metadane wykonania procedury, bez danych identyfikujących dziecko i bez opisu incydentu."),
            ("Bezpieczny rejestr", "Odrębny, ograniczony dostępem kanał dokumentowania interwencji, zatwierdzony przez właściciela po ocenie prawnej i prywatności."),
        ],
    )

    doc.add_heading("3. Role i odpowiedzialność", level=1)
    add_roles_table(doc)
    add_callout(
        doc,
        "Ciągłość dyżuru",
        "Dla każdego planowanego przyjazdu z dzieckiem musi być wiadomo, kto wykonuje procedurę i kto jest osiągalnym zastępcą. "
        "Jeżeli klucz wydaje członek rodziny lub inna osoba doraźna, przed pierwszym dyżurem przechodzi szkolenie i składa oświadczenie.",
    )

    doc.add_heading("4. Bezpieczne relacje personelu z dziećmi", level=1)
    doc.add_heading("4.1 Zachowania wymagane", level=2)
    add_bullets(
        doc,
        [
            "Przedstaw się, wyjaśniaj krótko, co robisz i dlaczego, oraz używaj języka odpowiedniego do wieku dziecka.",
            "Rozmawiaj w widocznym miejscu; gdy to możliwe, w obecności opiekuna, chyba że opiekun może stanowić zagrożenie.",
            "Szanuj granice, prywatność, odmowę kontaktu fizycznego i potrzebę dodatkowego czasu lub alternatywnej komunikacji.",
            "Kontakt fizyczny ogranicz do sytuacji koniecznych dla bezpieczeństwa lub pomocy i uprzedzaj o nim, jeśli sytuacja na to pozwala.",
            "Wejście do domku/pokoju ogranicz do uzasadnionych obowiązków, po zapowiedzi i zgodnie z zasadami obiektu; wyjątkiem jest nagłe zagrożenie.",
            "Informacje o dziecku przekazuj wyłącznie osobom, które muszą je znać dla ochrony lub wykonania obowiązku prawnego.",
        ],
        bullet_id,
    )
    doc.add_heading("4.2 Zachowania niedozwolone", level=2)
    add_bullets(
        doc,
        [
            "Przemoc, groźby, krzyk, poniżanie, zawstydzanie, dyskryminacja, komentarze seksualne lub żarty przekraczające granice.",
            "Dotykanie bez potrzeby, kontakt seksualny, proponowanie alkoholu, nikotyny, substancji odurzających albo niebezpiecznych aktywności.",
            "Prywatne wiadomości, zapraszanie do domu, spotkania poza obowiązkami, wymiana prezentów lub pieniędzy tworząca zależność.",
            "Fotografowanie, nagrywanie lub publikowanie wizerunku dziecka bez odrębnej legalnej podstawy i właściwej zgody.",
            "Pozostawanie sam na sam w zamkniętym pomieszczeniu bez uzasadnienia; ukrywanie kontaktu przed innymi członkami personelu.",
            "Proszenie dziecka o zachowanie tajemnicy dotyczącej zachowania personelu albo ignorowanie jego sprzeciwu i sygnałów dyskomfortu.",
        ],
        bullet_id,
    )
    add_callout(
        doc,
        "Kontakt online",
        "Personel używa wyłącznie zatwierdzonych kanałów obiektu i kontaktuje się z pełnoletnim rezerwującym. "
        "Nie nawiązuje prywatnych relacji z dzieckiem w mediach społecznościowych.",
        tone="warn",
    )

    doc.add_heading("5. Identyfikacja dziecka i relacji z dorosłym", level=1)
    p = doc.add_paragraph(
        "Procedurę wykonuje się dla każdego pobytu, w którym zgłoszono dziecko. Informację o obowiązku wyjaśnienia relacji "
        "przekazuje się przed przyjazdem. Przy bezobsługowym przekazaniu kodu procedurę przeprowadza się w zaplanowanej rozmowie "
        "na żywo lub osobiście przed udostępnieniem wejścia."
    )
    doc.add_heading("5.1 Zwykły przebieg", level=2)
    add_steps(
        doc,
        [
            "Sprawdź w Stawy OS, że pobyt wymaga procedury, oraz potwierdź właściwy domek i termin — bez wyświetlania danych dziecka.",
            "Wyjaśnij dorosłemu: „Stosujemy tę samą krótką procedurę przy każdym pobycie z dzieckiem, aby dbać o bezpieczeństwo najmłodszych gości”.",
            "Zapytaj dorosłego o relację z dzieckiem i obserwuj, czy kontakt jest swobodny, opiekuńczy i spójny z odpowiedziami.",
            "Jeżeli dorosły nie jest rodzicem lub opiekunem prawnym, poproś o wykazanie upoważnienia albo umożliwienie kontaktu z opiekunem. Oceniaj całość sytuacji, nie pojedynczy brak dokumentu.",
            "Gdy pozostają wątpliwości, porozmawiaj z dzieckiem spokojnie, językiem odpowiednim do wieku, najlepiej w widocznym miejscu i bez pytań sugerujących odpowiedź.",
            "Jeżeli relacja została racjonalnie wyjaśniona i nie ma sygnałów zagrożenia, w Stawy OS wybierz „Wykonano bez uwag”. Dopiero wtedy można wydać klucz lub kod.",
            "Jeżeli relacji nie można wyjaśnić albo pojawia się sygnał zagrożenia, nie wydawaj klucza/kodu; wybierz „Wymaga reakcji” i przejdź do rozdziału 7.",
        ],
        number_id,
    )
    doc.add_heading("5.2 Dokumenty i minimalizacja danych", level=2)
    add_bullets(
        doc,
        [
            "Pierwszym narzędziem są spokojna rozmowa i obserwacja; okazanie dokumentu jest środkiem pomocniczym stosowanym proporcjonalnie.",
            "Jeżeli gość dobrowolnie okazuje istniejący dokument lub upoważnienie, personel sprawdza go wzrokowo tylko w zakresie potrzebnym do wyjaśnienia relacji.",
            "Nie wykonuje się kopii, zdjęć ani skanów dokumentów dziecka lub dorosłego na potrzeby rutynowego wykonania standardu.",
            "Nie zapisuje się numeru dokumentu, PESEL-u, adresu dziecka, wizerunku ani treści dokumentu w Stawy OS.",
            "Odmowa okazania dokumentu nie jest samodzielną podstawą oskarżenia; może jednak — wraz z innymi niespójnościami — uzasadniać wstrzymanie dostępu i kontakt z koordynatorem lub Policją.",
        ],
        bullet_id,
    )
    add_callout(
        doc,
        "Zakaz obchodzenia procedury",
        "Presja czasu, późny przyjazd, zdenerwowanie gościa, opłacona rezerwacja ani znajomość z właścicielem nie są podstawą pominięcia procedury.",
        tone="danger",
    )

    doc.add_heading("6. Sygnały wymagające uwagi", level=1)
    add_bullets(
        doc,
        [
            "Dziecko sygnalizuje strach, prosi o pomoc, mówi o przemocy albo nie chce zostać z dorosłym.",
            "Odpowiedzi dorosłego i dziecka są wyraźnie niespójne, wyuczone lub dorosły uniemożliwia dziecku wypowiedź.",
            "Relacja nie wygląda na swobodną i opiekuńczą; pojawiają się groźby, kontrola, poniżanie, przemoc lub seksualizujące zachowania.",
            "Widoczne obrażenia, silne zaniedbanie, brak podstawowej opieki, nieadekwatny ubiór lub zachowanie mogące wskazywać na odurzenie.",
            "Nietypowa prośba o pełną anonimowość, częste zmiany dorosłych, brak logicznego celu podróży albo próba izolowania dziecka.",
            "Materiały, przedmioty lub sytuacja w domku wskazują na możliwość seksualnego wykorzystania, handlu ludźmi lub innego przestępstwa.",
        ],
        bullet_id,
    )
    add_callout(
        doc,
        "Sygnał to nie wyrok",
        "Pojedyncza cecha nie przesądza o krzywdzeniu. Personel zapisuje fakty, nie diagnozy, i ocenia pilność ochrony, nie winę.",
        tone="info",
    )

    doc.add_heading("7. Procedura reakcji", level=1)
    doc.add_heading("7.1 Bezpośrednie zagrożenie", level=2)
    add_steps(
        doc,
        [
            "Zadzwoń 112; w razie potrzeby udziel pierwszej pomocy w granicach swoich umiejętności.",
            "Przenieś dziecko do bezpiecznego, widocznego miejsca. Nie pozostawiaj go z osobą stanowiącą możliwe zagrożenie.",
            "Nie informuj osoby podejrzewanej o działaniach, jeżeli może to zwiększyć ryzyko, utrudnić interwencję lub zniszczyć dowody.",
            "Wykonuj polecenia służb. Nie stosuj przymusu fizycznego poza konieczną obroną życia lub zdrowia zgodnie z prawem.",
            "Powiadom koordynatora i zabezpiecz informacje, nagrania lub przedmioty wyłącznie w zakresie wskazanym przez służby.",
        ],
        number_id,
    )
    doc.add_heading("7.2 Podejrzenie bez bezpośredniego zagrożenia", level=2)
    add_steps(
        doc,
        [
            "Wstrzymaj wydanie klucza lub kodu i zapewnij spokojne warunki.",
            "Powiadom koordynatora; jeżeli jest nieosiągalny, skontaktuj się z zastępcą, dyżurnym KPP Międzyrzecz lub 112.",
            "Koordynator ocenia właściwy kanał: Policja/prokuratura przy podejrzeniu przestępstwa, sąd opiekuńczy przy zdarzeniu uzasadniającym postępowanie z urzędu, GOPS przy potrzebie wsparcia rodziny. Pilność ochrony ma pierwszeństwo przed konsultacją.",
            "Nie prowadź konfrontacji i nie próbuj samodzielnie potwierdzać wersji zdarzeń. Przekaż służbom obserwacje i spontaniczne wypowiedzi.",
            "Zapisz interwencję w bezpiecznym rejestrze i numer referencyjny w Stawy OS. Dalsze działania prowadź według instrukcji właściwego organu.",
        ],
        number_id,
    )
    doc.add_heading("7.3 Rozmowa z dzieckiem", level=2)
    add_bullets(
        doc,
        [
            "Zachowaj spokój; powiedz: „Dziękuję, że mi o tym mówisz. To nie jest twoja wina. Postaram się zapewnić ci bezpieczeństwo”.",
            "Nie obiecuj pełnej tajemnicy: wyjaśnij, że informację przekażesz tylko osobom potrzebnym do pomocy.",
            "Używaj pytań otwartych i koniecznych: „Co się stało?”, „Czego teraz potrzebujesz?”, „Czy teraz jesteś bezpieczny/bezpieczna?”.",
            "Nie pytaj „dlaczego”, nie sugeruj sprawcy ani szczegółów, nie proś o wielokrotne powtarzanie i nie oceniaj wiarygodności.",
            "Jeżeli dziecko komunikuje się inaczej, zapewnij czas, prosty język, możliwość pisania, wskazywania albo wsparcia dostępnej osoby — bez udziału osoby podejrzewanej.",
        ],
        bullet_id,
    )

    doc.add_heading("8. Dokumentowanie i ochrona danych", level=1)
    add_callout(
        doc,
        "Dwa oddzielne miejsca zapisu",
        "Stawy OS przechowuje wyłącznie wykonanie procedury i numer sprawy. Dane identyfikujące dziecko, opis, dowody i korespondencja ze służbami trafiają tylko do bezpiecznego rejestru interwencji.",
        tone="info",
    )
    doc.add_heading("8.1 Rutynowe wykonanie", level=2)
    add_bullets(
        doc,
        [
            "Stawy OS: identyfikator pobytu, wersja SOP, czas, operator i wynik „Bez uwag” albo „Wymaga reakcji”.",
            "Brak pól na imię, nazwisko, datę urodzenia, PESEL, numer dokumentu, zdjęcie, skan lub opis zdarzenia.",
            "Rutynowy zapis bez uwag przechowuje się 24 miesiące od zakończenia pobytu, a następnie usuwa lub anonimizuje w kontrolowanym procesie.",
        ],
        bullet_id,
    )
    doc.add_heading("8.2 Interwencja", level=2)
    add_bullets(
        doc,
        [
            "Karta zawiera tylko konieczne fakty: data, miejsce, osoba zgłaszająca, źródło informacji, zaobserwowane fakty, spontaniczne słowa dziecka, działania, kontakty ze służbami i wynik przekazania.",
            "Dane osób oraz materiały dowodowe zapisuje się tylko wtedy, gdy są potrzebne do ochrony, zgłoszenia lub współpracy z organem.",
            "Dostęp mają wyłącznie formalnie upoważnione osoby. Każdy dostęp i przekazanie powinny być możliwe do odtworzenia.",
            "Pliki szyfruje się w spoczynku i podczas przesyłania; zabronione są prywatne telefony, komunikatory, zwykły e-mail i niezabezpieczone nośniki.",
            "Domyślny okres retencji projektu: 6 lat od końca roku zamknięcia sprawy; bieg wstrzymuje toczące się postępowanie, obowiązek prawny lub legal hold. Właściciel zatwierdza tę decyzję po konsultacji prawnej/RODO.",
            "Co najmniej raz w roku koordynator przegląda podstawę, zakres i dalszą potrzebę przechowywania. Usunięcie jest dokumentowane.",
        ],
        bullet_id,
    )
    doc.add_heading("8.3 Nagrania i dowody", level=2)
    add_bullets(
        doc,
        [
            "Nie zmieniaj, nie opisuj na oryginale i nie rozpowszechniaj materiału. Zapisz, kto i kiedy go zabezpieczył.",
            "Jeżeli monitoring może obejmować zdarzenie, koordynator zabezpiecza właściwy zakres czasowy przed automatycznym nadpisaniem.",
            "Materiał przekazuje się właściwemu organowi bez zbędnych kopii. Publiczne udostępnienie lub wysłanie w grupie personelu jest zabronione.",
        ],
        bullet_id,
    )

    doc.add_heading("9. Personel, weryfikacja i przygotowanie", level=1)
    add_bullets(
        doc,
        [
            "Każda osoba przed samodzielnym kontaktem operacyjnym z gośćmi zapoznaje się ze standardem, przechodzi instruktaż scenariuszowy i podpisuje oświadczenie.",
            "Szkolenie obejmuje: sześć kroków przed kluczem, sygnały ryzyka, rozmowę z dzieckiem, 112, dokumentowanie, prywatność i ćwiczenie na telefonie używanym przy przyjeździe.",
            "Osoby faktycznie wykonujące działalność objętą art. 21 ustawy — np. opiekę, organizowanie wypoczynku, sportu lub zajęć dla dzieci — są weryfikowane w wymaganych rejestrach i dokumentach przed dopuszczeniem. Sama zwykła obsługa noclegu wymaga odrębnej klasyfikacji obowiązków.",
            "Po zmianie standardu każda osoba otrzymuje krótkie szkolenie aktualizujące. Raz w roku wykonuje się próbę scenariusza, a po incydencie — omówienie procesu bez ujawniania danych ponad konieczny zakres.",
            "Osoba nieprzygotowana nie wydaje kluczy/kodów dla pobytów z dziećmi i przekazuje zadanie osobie przeszkolonej.",
        ],
        bullet_id,
    )

    doc.add_heading("10. Publikacja, dostępność i informowanie gości", level=1)
    add_bullets(
        doc,
        [
            "Pełną i skróconą wersję publikuje się na stronie internetowej oraz wywiesza w widocznym miejscu obiektu.",
            "Wersję dla dzieci umieszcza się na wysokości i w miejscu dla nich dostępnym, z dużym tekstem, wysokim kontrastem i numerami pomocy.",
            "Przed przyjazdem gość otrzymuje krótką informację, że procedura jest wykonywana przy każdym pobycie z dzieckiem.",
            "Na życzenie zapewnia się alternatywny sposób przekazania informacji: odczytanie, większy druk, prosty język, możliwość pisania lub komunikację wspomagającą.",
            "Wersje językowe mogą być przygotowane po polsku, angielsku i niemiecku, ale wersją zatwierdzoną pozostaje dokument wskazany w zarządzeniu.",
        ],
        bullet_id,
    )

    doc.add_heading("11. Przegląd, zgłoszenia i doskonalenie", level=1)
    add_bullets(
        doc,
        [
            "Każdy może zgłosić naruszenie koordynatorowi, właścicielowi lub bezpośrednio właściwym służbom; zgłoszenie nie może powodować odwetu.",
            "Koordynator prowadzi rejestr naruszeń i propozycji zmian bez ujawniania danych w analizach zbiorczych.",
            "Pierwsza ocena następuje po 6 miesiącach pilotażu; kolejna co najmniej raz na 2 lata oraz niezwłocznie po poważnym incydencie, zmianie procesu, danych lub prawa.",
            "Ocena obejmuje skuteczność procedury, czasy reakcji, wykonanie przed kluczem, dostępność dla dzieci, zakres danych, retencję, szkolenia i kontakt do służb.",
            "Wnioski dokumentuje się pisemnie. Nowa wersja otrzymuje numer, datę zatwierdzenia, datę obowiązywania, termin przeglądu i historię zmian.",
        ],
        bullet_id,
    )

    doc.add_heading("Załącznik A. Sześć kroków dla osoby wydającej klucz lub kod", level=1)
    add_steps(
        doc,
        [
            "OTWÓRZ właściwy pobyt w Stawy OS i potwierdź, że procedura ma status „Do wykonania”.",
            "WYJAŚNIJ dorosłemu, że standard dotyczy każdego pobytu z dzieckiem i służy bezpieczeństwu.",
            "USTAL relację dorosły–dziecko przez spokojne pytanie i obserwację; w razie potrzeby poproś o wykazanie upoważnienia lub kontakt z opiekunem.",
            "SPRAWDŹ SPÓJNOŚĆ. Gdy trzeba, porozmawiaj krótko z dzieckiem odpowiednio do jego wieku. Nie kopiuj dokumentów i nie zapisuj danych dziecka w aplikacji.",
            "ZAPISZ „Wykonano bez uwag” tylko wtedy, gdy relacja została racjonalnie wyjaśniona i nie ma sygnałów zagrożenia.",
            "ZATRZYMAJ WYDANIE KLUCZA/KODU i wybierz „Wymaga reakcji”, jeżeli sytuacja pozostaje niewyjaśniona lub niepokojąca. W zagrożeniu dzwoń 112; inaczej natychmiast eskaluj do koordynatora.",
        ],
        number_id,
    )
    add_callout(
        doc,
        "W aplikacji",
        "Nie wpisuj imienia dziecka, wieku, dokumentu, opisu zdarzenia ani podejrzeń. Po interwencji koordynator zamyka reakcję wyłącznie numerem wpisu w bezpiecznym rejestrze.",
        tone="danger",
    )

    doc.add_heading("Załącznik B. Gotowe komunikaty", level=1)
    scripts = [
        ("Na początku", "„W naszym obiekcie przy każdym pobycie z dzieckiem stosujemy krótką procedurę bezpieczeństwa. Zapytam o relację z dzieckiem; zwykle trwa to chwilę”."),
        ("Gdy dorosły nie jest opiekunem", "„Proszę pokazać upoważnienie albo umożliwić kontakt z rodzicem lub opiekunem. Nie wykonujemy kopii dokumentu”."),
        ("Gdy trzeba porozmawiać z dzieckiem", "„Cześć, mam na imię … i pomagam przy przyjeździe. Kim jest osoba, z którą przyjechałeś/przyjechałaś? Czy czujesz się teraz bezpiecznie?”."),
        ("Gdy wszystko jest w porządku", "„Dziękuję za chwilę. Procedura jest zakończona. Życzymy spokojnego pobytu”."),
        ("Gdy sytuacja pozostaje niewyjaśniona", "„Nie mogę jeszcze wydać klucza ani kodu. Muszę skontaktować się z osobą odpowiedzialną za bezpieczeństwo dzieci”."),
        ("Gdy dziecko coś ujawnia", "„Dziękuję, że mi o tym mówisz. To nie jest twoja wina. Nie mogę obiecać, że nikomu nie powiem, ale przekażę to tylko osobom, które mogą pomóc”."),
    ]
    for title, body in scripts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(title), size=11, color=GREEN, bold=True)
        p = doc.add_paragraph(body)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(7)

    doc.add_heading("Załącznik C. Karta interwencji — bezpieczny rejestr", level=1)
    add_callout(
        doc,
        "Nie przechowuj tej karty w Stawy OS",
        "Karta może zawierać dane szczególnej wrażliwości. Wymaga osobnego, zatwierdzonego miejsca, upoważnień i rejestru dostępu.",
        tone="danger",
    )
    fields = [
        "Numer sprawy nadany przez rejestr",
        "Data, godzina i miejsce zdarzenia / zgłoszenia",
        "Osoba sporządzająca i źródło informacji",
        "Fakty zaobserwowane bez interpretacji",
        "Spontaniczne słowa dziecka zapisane możliwie dokładnie",
        "Ocena, czy występowało bezpośrednie zagrożenie",
        "Działania ochronne i pierwsza pomoc",
        "Osoby oraz instytucje powiadomione, czas, kanał i osoba przyjmująca",
        "Przekazane materiały i ciąg dostępu do dowodów",
        "Decyzje służb / dalsze polecenia",
        "Data zamknięcia, podstawa i termin kolejnego przeglądu retencji",
    ]
    for field in fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(f"{field}: "), size=10.5, color=INK, bold=True)
        set_run_font(p.add_run("................................................................................................"), size=10, color=MUTED)

    doc.add_heading("Załącznik D. Oświadczenie i szkolenie personelu", level=1)
    p = doc.add_paragraph(
        "Oświadczam, że zapoznałem/am się z obowiązującą wersją Standardów ochrony małoletnich Stawy u Sikory, "
        "rozumiem sześć kroków przed wydaniem klucza/kodu, zasady eskalacji i zakaz zapisywania danych dziecka w Stawy OS. "
        "Zobowiązuję się stosować Standardy oraz zgłaszać naruszenia."
    )
    add_key_value_table(
        doc,
        [
            ("Imię i nazwisko", "................................................................................................"),
            ("Rola", "................................................................................................"),
            ("Wersja standardu", "................................................................................................"),
            ("Forma szkolenia", "instruktaż / scenariusz / ćwiczenie w aplikacji / aktualizacja"),
            ("Data i podpis", "................................................................................................"),
            ("Osoba szkoląca", "................................................................................................"),
        ],
    )

    doc.add_heading("Załącznik E. Kontrola zatwierdzenia i przeglądu", level=1)
    checks = [
        "Potwierdzono formalną nazwę podmiotu i zakres obiektów.",
        "Imiennie powołano koordynatora i zastępcę oraz opublikowano sposób kontaktu.",
        "Prawnik lub inna uprawniona osoba przeprowadziła przegląd zgodności dokumentu.",
        "Administrator danych zatwierdził bezpieczny rejestr, upoważnienia, retencję i procedurę usuwania.",
        "Przeszkolono wszystkie osoby wydające klucze/kody, w tym członków rodziny i zastępstwa.",
        "Opublikowano pełną i dziecięcą wersję na stronie oraz wywieszono obie w obiekcie.",
        "Sprawdzono numery kontaktowe, działanie 112 i osiągalność koordynatora.",
        "Przeprowadzono próbę: zwykły przyjazd, brak dokumentu, osoba niebędąca rodzicem, ujawnienie przemocy, późny przyjazd i niedostępny koordynator.",
        "W Stawy OS aktywowano właściwą wersję, datę obowiązywania, termin przeglądu, linki i sześć kroków.",
        "Po przeglądzie sporządzono pisemne wnioski, decyzje i historię zmian.",
    ]
    add_bullets(doc, checks, bullet_id)
    add_key_value_table(
        doc,
        [
            ("Zatwierdził/a", "................................................................................................"),
            ("Funkcja", "................................................................................................"),
            ("Data wejścia w życie", "................................................................................................"),
            ("Termin przeglądu (maks. 2 lata)", "................................................................................................"),
            ("Podpis", "................................................................................................"),
        ],
    )

    doc.add_heading("Załącznik F. Kontakty operacyjne", level=1)
    add_contact_table(doc)
    add_callout(
        doc,
        "Aktualizacja",
        "Koordynator sprawdza tabelę kontaktów przy każdym przeglądzie i co najmniej raz w roku. W sytuacji pilnej właściwy jest numer 112.",
        tone="warn",
    )
    add_sources(doc)
    return doc


def build_child_document():
    doc = Document()
    configure_document(doc, child=True)
    bullet_id = add_numbering_definition(doc, bullet=True)
    number_id = add_numbering_definition(doc, bullet=False)
    add_title_page(doc, child=True)

    doc.add_heading("Masz tutaj prawa", level=1)
    add_bullets(
        doc,
        [
            "Masz prawo do szacunku, spokoju i prywatności.",
            "Nikt nie może cię bić, straszyć, poniżać ani dotykać w sposób, którego nie chcesz.",
            "Nikt nie może prosić cię o seksualne zdjęcia, wiadomości ani zachowanie krzywdzącej tajemnicy.",
            "Możesz powiedzieć „nie”, odejść w widoczne miejsce i poprosić o pomoc.",
            "Jeśli potrzebujesz więcej czasu, prostszych słów, pisania albo innego sposobu rozmowy — powiedz nam.",
        ],
        bullet_id,
        child=True,
    )

    doc.add_heading("Dlaczego możemy zadać pytanie?", level=1)
    p = doc.add_paragraph(
        "Przy pobycie z osobą poniżej 18 lat pracownik może zapytać dorosłego, kim jest dla dziecka. "
        "Czasem zapyta też ciebie. Robimy to przy każdym takim pobycie, żeby dzieci były bezpieczne."
    )
    set_run_font(p.runs[0], name="Arial", size=13.5)
    add_callout(
        doc,
        "To nie jest twoja wina",
        "Jeśli dorosły się złości albo sytuacja jest niejasna, nie odpowiadasz za jego reakcję. Możesz powiedzieć tylko tyle, ile potrafisz.",
        tone="info",
        child=True,
    )

    doc.add_heading("Jeśli coś cię niepokoi", level=1)
    add_steps(
        doc,
        [
            "Przejdź do widocznego miejsca albo podejdź do pracownika.",
            "Powiedz: „Nie czuję się bezpiecznie” albo „Potrzebuję pomocy”. Możesz też napisać lub pokazać tę stronę.",
            "Jeśli możesz, powiedz, czego potrzebujesz teraz. Nie musisz opowiadać wszystkiego kilka razy.",
            "W nagłym niebezpieczeństwie dzwoń 112. Możesz też zadzwonić do zaufanej osoby.",
        ],
        number_id,
        child=True,
    )
    add_callout(
        doc,
        "Ważne",
        "Pracownik cię wysłucha. Nie może obiecać, że zachowa wszystko w tajemnicy, bo czasem musi wezwać osoby, które mogą cię ochronić. Informację powinien przekazać tylko tym osobom.",
        tone="warn",
        child=True,
    )
    doc.add_heading("Gdzie dostać pomoc", level=1)
    emergency = doc.add_paragraph()
    emergency.alignment = WD_ALIGN_PARAGRAPH.CENTER
    emergency.paragraph_format.space_after = Pt(12)
    set_run_font(emergency.add_run("112"), name="Arial", size=34, color=RED, bold=True)
    p = doc.add_paragraph("Numer alarmowy — gdy ktoś jest teraz w niebezpieczeństwie.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.runs[0], name="Arial", size=15, color=BLACK, bold=True)

    help_items = [
        ("116 111", "Telefon zaufania dla dzieci i młodzieży — bezpłatnie, całą dobę."),
        ("800 12 12 12", "Dziecięcy Telefon Zaufania Rzecznika Praw Dziecka — bezpłatnie, całą dobę."),
        ("czat.brpd.gov.pl", "Bezpłatny i anonimowy czat z konsultantem — całą dobę."),
        ("501 510 005", "Kontakt Stawy u Sikory — noclegi."),
    ]
    for idx, (number, description) in enumerate(help_items):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        set_paragraph_box(
            p,
            color="BFD2C9",
            fill=LIGHT_GREEN if idx % 2 == 0 else LAKE,
            size="8",
        )
        set_run_font(p.add_run(number), name="Arial", size=20, color=GREEN, bold=True)
        run = p.add_run()
        run.add_break()
        set_run_font(p.add_run(description), name="Arial", size=13, color=BLACK)

    doc.add_heading("Możesz powiedzieć", level=1)
    phrases = [
        "„Nie chcę zostać sam/sama z tą osobą”.",
        "„Proszę zadzwoń po pomoc”.",
        "„Wolę napisać niż mówić”.",
        "„Chcę porozmawiać bez tej osoby obok”.",
    ]
    add_bullets(doc, phrases, bullet_id, child=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    set_run_font(
        p.add_run("Pamiętaj: przemoc nigdy nie jest twoją winą. Proś o pomoc tyle razy, ile potrzebujesz."),
        name="Arial",
        size=13,
        color=RED,
        bold=True,
    )
    return doc


def save_documents():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    full = build_full_document()
    child = build_child_document()
    full.core_properties.title = "Standardy ochrony małoletnich — Stawy u Sikory — projekt 0.9"
    full.core_properties.subject = "SOP dla obiektu świadczącego usługi noclegowe"
    full.core_properties.keywords = "ochrona małoletnich, SOP, hotel, noclegi"
    child.core_properties.title = "Ochrona małoletnich — wersja dla dzieci — Stawy u Sikory — projekt 0.9"
    child.core_properties.subject = "Skrócona wersja standardów dla dzieci i młodzieży"
    full.save(FULL_PATH)
    child.save(CHILD_PATH)
    print(FULL_PATH)
    print(CHILD_PATH)


if __name__ == "__main__":
    save_documents()
